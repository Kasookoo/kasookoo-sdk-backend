from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Right-click and Update Field to generate table of contents."
    fld_sep.append(txt)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    doc.add_page_break()


def render_markdown_to_docx(markdown_path: Path, docx_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    add_toc(doc)

    in_code = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            if not in_code:
                doc.add_paragraph("")
            continue

        if in_code:
            doc.add_paragraph(line, style="No Spacing")
            continue

        if not stripped:
            doc.add_paragraph("")
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif len(stripped) > 3 and stripped[0].isdigit() and stripped[1] == "." and stripped[2] == " ":
            doc.add_paragraph(stripped[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)


if __name__ == "__main__":
    pairs = [
        ("BUSINESS_REQUIREMENTS_DOCUMENT.md", "BUSINESS_REQUIREMENTS_DOCUMENT.docx"),
        ("SDK_DEVELOPER_IMPLEMENTATION_GUIDE.md", "SDK_DEVELOPER_IMPLEMENTATION_GUIDE.docx"),
    ]
    root = Path(__file__).resolve().parents[1]
    for md_file, docx_file in pairs:
        render_markdown_to_docx(root / md_file, root / docx_file)
    print("Updated both documents with TOC.")
